"""
scripts/md_to_docx.py — Gera .docx a partir de um documento de `docs/` (para quem vai EDITAR).

Complementa `docs_to_pdf.py`: o PDF é para ler e arquivar; o .docx é para a orientadora, a
assessoria ou o CEP **editarem e comentarem**. Por isso monta estilos NATIVOS do Word (Heading,
List Bullet, tabelas de verdade) em vez de converter HTML — conversão produz marcação que trava
o controle de alterações, e aí o formato editável não serve para o que existe.

Mesma fonte única do PDF: o `.md` do repositório. Não há segunda redação em lugar nenhum.

Uso:
    python scripts/md_to_docx.py pendencias
    python scripts/md_to_docx.py tcle --out-dir "C:/caminho"
    python scripts/md_to_docx.py --list
"""
from __future__ import annotations
import argparse
import pathlib
import re
import sys
from html.parser import HTMLParser

try:
    import markdown
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:                                      # noqa: BLE001
    sys.exit("Faltam dependências: pip install markdown python-docx")

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
from docs_to_pdf import DATE, DOCS, DOCUMENTOS, ROOT     # noqa: E402  (mesmo catálogo)

TEAL = RGBColor(0x0F, 0x5C, 0x68)
PETROL = RGBColor(0x1B, 0x4B, 0x5A)
INK = RGBColor(0x1A, 0x26, 0x32)
MUTED = RGBColor(0x5A, 0x6B, 0x78)
ALERT = RGBColor(0xB3, 0x5A, 0x12)


def shade_cell(cell, hexcolor: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)


class DocxBuilder(HTMLParser):
    """Percorre o HTML do Markdown e emite elementos nativos do Word.

    Só trata o subconjunto que os documentos de `docs/` usam (títulos, parágrafos, listas,
    tabelas, citações, régua). Marcação fora disso vira texto simples — de propósito: é melhor
    perder um detalhe visual do que gerar um .docx que não abre direito."""

    INLINE = {"strong", "b", "em", "i", "code", "a"}

    def __init__(self, doc: Document) -> None:
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.para = None            # parágrafo corrente
        self.fmt: list[str] = []    # pilha de formatação inline
        self.list_kind: list[str] = []
        self.table = None
        self.row = None
        self.cell_idx = 0
        self.in_head = False
        self.quote_cell = None      # célula da "caixa" de citação
        self.pending_cols = 0

    # -------- utilidades --------
    def _new_para(self, style: str | None = None, container=None):
        host = container if container is not None else self.doc
        p = host.add_paragraph(style=style) if style else host.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        return p

    def _emit(self, text: str) -> None:
        if self.para is None or not text:
            return
        run = self.para.add_run(text)
        run.bold = any(f in ("strong", "b") for f in self.fmt)
        run.italic = any(f in ("em", "i") for f in self.fmt)
        if "code" in self.fmt:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED

    # -------- tags --------
    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        if tag in self.INLINE:
            self.fmt.append(tag)
        elif tag in ("h1", "h2", "h3"):
            lvl = int(tag[1])
            h = self.doc.add_heading(level=min(lvl, 3))
            h.paragraph_format.space_before = Pt(14 if lvl <= 2 else 10)
            h.paragraph_format.space_after = Pt(4)
            self.para = h
            self._style_next = (TEAL if lvl <= 2 else PETROL, 15 if lvl == 1 else (12 if lvl == 2 else 11))
        elif tag == "p":
            host = self.quote_cell if self.quote_cell is not None else None
            self.para = self._new_para(container=host)
            if host is None:
                self.para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif tag in ("ul", "ol"):
            self.list_kind.append("List Bullet" if tag == "ul" else "List Number")
        elif tag == "li":
            style = self.list_kind[-1] if self.list_kind else "List Bullet"
            self.para = self._new_para(style=style)
        elif tag == "blockquote":
            # Caixa de destaque = tabela de 1 célula com fundo (renderiza em qualquer Word).
            t = self.doc.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            self.quote_cell = t.rows[0].cells[0]
            self.quote_cell.paragraphs[0]._p.getparent().remove(self.quote_cell.paragraphs[0]._p)
            self._quote_alert = False
        elif tag == "table":
            self.table = self.doc.add_table(rows=0, cols=0)
            self.table.style = "Table Grid"
            self.pending_cols = 0
        elif tag == "thead":
            self.in_head = True
        elif tag == "tr":
            if self.table is not None:
                if not self.table.columns:
                    self.row = None                     # colunas criadas no 1º th/td
                else:
                    self.row = self.table.add_row()
                self.cell_idx = 0
        elif tag in ("th", "td"):
            if self.table is None:
                return
            if self.row is None:                        # primeira linha: cria colunas sob demanda
                self.table.add_column(Cm(2))
                if len(self.table.rows) == 0:
                    self.table.add_row()
                self.row = self.table.rows[0]
            cell = self.row.cells[min(self.cell_idx, len(self.row.cells) - 1)]
            if self.in_head:
                shade_cell(cell, "EEF3F5")
            self.para = cell.paragraphs[0]
            self.para.paragraph_format.space_after = Pt(1)
        elif tag == "hr":
            p = self._new_para()
            p.add_run("_" * 60).font.color.rgb = RGBColor(0xCF, 0xDB, 0xE2)
            self.para = None

    def handle_endtag(self, tag):
        if tag in self.INLINE:
            if self.fmt and self.fmt[-1] == tag:
                self.fmt.pop()
        elif tag in ("h1", "h2", "h3"):
            color, size = getattr(self, "_style_next", (TEAL, 12))
            for r in self.para.runs:
                r.font.color.rgb = color
                r.font.size = Pt(size)
                r.bold = True
            self.para = None
        elif tag in ("p", "li"):
            self.para = None
        elif tag in ("ul", "ol"):
            if self.list_kind:
                self.list_kind.pop()
        elif tag == "blockquote":
            if self.quote_cell is not None:
                shade_cell(self.quote_cell, "FDF4EC" if self._quote_alert else "F2F6F8")
                self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
            self.quote_cell = None
            self.para = None
        elif tag == "thead":
            self.in_head = False
        elif tag in ("th", "td"):
            self.cell_idx += 1
            self.para = None
        elif tag == "table":
            self.table = self.row = None

    def handle_data(self, data):
        if self.para is None:
            return
        text = data if "code" in self.fmt else re.sub(r"\s+", " ", data)
        if not text.strip() and self.para.runs:
            text = " "
        elif not text.strip():
            return
        if self.quote_cell is not None and "⚠️" in text:
            self._quote_alert = True
            text = text.replace("⚠️", "").lstrip()
        self._emit(text)


def build(key: str, out_dir: pathlib.Path) -> pathlib.Path:
    src_name, out_name, band_t, band_p, ident = DOCUMENTOS[key]
    md_text = (DOCS / src_name).read_text(encoding="utf-8")
    html = markdown.markdown(md_text, extensions=["tables", "sane_lists", "fenced_code"])

    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.left_margin = s.right_margin = Cm(1.9)
    s.top_margin, s.bottom_margin = Cm(1.9), Cm(1.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    # Tarja de status (a mesma do PDF) — sem ela o .docx circularia sem dizer que é rascunho.
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, "FDF4EC")
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    r = p0.add_run(band_t.format(date=DATE))
    r.bold, r.font.size, r.font.color.rgb = True, Pt(9), ALERT
    p1 = cell.add_paragraph()
    r1 = p1.add_run(re.sub(r"<[^>]+>", "", band_p))
    r1.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    DocxBuilder(doc).feed(html)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(14)
    rf = foot.add_run(ident.format(date=DATE))
    rf.font.size, rf.font.color.rgb = Pt(8), MUTED

    dest = out_dir / f"{out_name}_{DATE}.docx"
    doc.save(dest)
    return dest


def main() -> int:
    ap = argparse.ArgumentParser(description="Gera .docx editável de um documento de docs/.")
    ap.add_argument("docs", nargs="*", help="chaves (padrão: pendencias)")
    ap.add_argument("--list", action="store_true")
    ap.add_argument("--out-dir", default=str(ROOT.parent))
    args = ap.parse_args()

    if args.list:
        for k, (src, *_ ) in DOCUMENTOS.items():
            print(f"  {k:12} {src}")
        return 0

    chaves = args.docs or ["pendencias"]
    if desconhecidas := [k for k in chaves if k not in DOCUMENTOS]:
        sys.exit(f"chave(s) desconhecida(s): {', '.join(desconhecidas)}")

    out_dir = pathlib.Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    for k in chaves:
        dest = build(k, out_dir)
        print(f"  [ok] {k:12} -> {dest.name}  ({dest.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
